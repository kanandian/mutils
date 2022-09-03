# import docx
from docx import Document
import xlwt
from collections import defaultdict
import re
import os

all_items = ['名称','时间','地点']
selected_items = ['名称','时间','地点']
separators = ['，','：','；']
folder_path = '/Users/macpro/PycharmProjects/mutils/tmp/target'

def write_to_excel(data):
    global item_cnt
    for idx, item in enumerate(selected_items):
        worksheet.write(item_cnt, idx, ','.join(data[item]))
    item_cnt+=1


def handle_file(file_path):
    print(file_path)
    data = defaultdict(list)
    doc = Document(file_path)
    for idx, para in enumerate(doc.paragraphs):
        # print(idx, para.text)
        txt = para.text
        parts = re.split(separators, txt)
        item_name = None
        for part in parts:
            if part in all_items:
                item_name = part
            elif item_name is not None:
                data[item_name].append(part)
    write_to_excel(data)

item_cnt=1
separators = '|'.join(separators)
files = os.listdir(folder_path)
workbook = xlwt.Workbook(encoding='utf8')
worksheet = workbook.add_sheet('1')
for idx, item in enumerate(selected_items):
    worksheet.write(0, idx, item)
for file_name in files:
    file_path = folder_path+'/'+file_name
    # print(file_name)
    # print(file_path)
    handle_file(file_path)
workbook.save('output1.xls')

# from win32com import client
# def doc2docx(doc_name, docx_name):
#     try:
#         # 首先将doc转换成docx
#         word = client.Dispatch("Word.Application")
#         doc = word.Documents.Open(doc_name)
#         # 使用参数16表示将doc转换成docx
#         doc.SaveAs(docx_name, 16)
#         doc.Close()
#         word.Quit()
#     except:
#         pass


# workbook = xlwt.Workbook(encoding='utf8')
# worksheet = workbook.add_sheet('1')
#
# doc = Document(doc_path)
# cnt=0
# for para in doc.paragraphs:
#     cnt+=1
#     content = para.text
#     print(content+' '+str(cnt))
#     if content=='' or content is None:
#         continue
#     worksheet.write(cnt, 1, content)
#
# workbook.save('output1.xls')